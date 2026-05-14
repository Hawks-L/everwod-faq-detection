"""Convierte GUIA_DEMOSTRACION.md a GUIA_DEMOSTRACION.docx con formato profesional."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "GUIA_DEMOSTRACION.md"
DST = "GUIA_DEMOSTRACION.docx"

doc = Document()

section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def add_table_from_md(doc, lines):
    rows = [l for l in lines if l.startswith("|") and "---" not in l]
    if not rows:
        return
    cells = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            if i == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "1F3864")
                tcPr.append(shd)

with open(SRC, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if line.startswith("# ") and not line.startswith("## "):
        p = doc.add_heading(line[2:], level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1

    elif line.startswith("## ") and not line.startswith("### "):
        doc.add_heading(line[3:], level=2)
        i += 1

    elif line.startswith("### "):
        doc.add_heading(line[4:], level=3)
        i += 1

    elif line.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].rstrip())
            i += 1
        add_table_from_md(doc, table_lines)
        doc.add_paragraph()

    elif line.startswith("```"):
        code_lines = []
        i += 1
        while i < len(lines) and not lines[i].startswith("```"):
            code_lines.append(lines[i].rstrip())
            i += 1
        i += 1
        p = doc.add_paragraph()
        p.style = "No Spacing"
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    elif line.startswith("---"):
        doc.add_paragraph().add_run("─" * 80).font.color.rgb = RGBColor(180, 180, 180)
        i += 1

    elif line.startswith("- ") or line.startswith("• "):
        text = line[2:].strip()
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        doc.add_paragraph(text, style="List Bullet")
        i += 1

    elif line == "":
        doc.add_paragraph()
        i += 1

    elif re.match(r"^\*\*.+?\*\*:", line):
        p = doc.add_paragraph()
        for part in re.split(r"(\*\*.+?\*\*)", line):
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(re.sub(r"`(.+?)`", r"\1", part))
        i += 1

    else:
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        text = re.sub(r"`(.+?)`", r"\1", text)
        if text.strip():
            doc.add_paragraph(text)
        i += 1

doc.save(DST)
print(f"Documento guardado: {DST}")
